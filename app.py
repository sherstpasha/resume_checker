import os
import gradio as gr
import requests
from dotenv import load_dotenv
from docx import Document
import textract
from pypdf import PdfReader
import json
import re

# Загружаем переменные окружения
load_dotenv()
API_BASE_URL = os.getenv("API_BASE_URL", "http://127.0.0.1:8080/v1").rstrip("/")
MODEL_NAME = os.getenv("MODEL_NAME", "llama-2-7b-chat")
JOB_DESCRIPTION_PATH = os.getenv("JOB_DESCRIPTION_PATH")

# ---------- Очистка reasoning ----------
def strip_thinking_tags(text: str, tags=None) -> str:
    final_marker = "<|channel|>final<|message|>"
    if final_marker in text:
        text = text.split(final_marker, 1)[1]

    if tags is None:
        tags = ["think", "reflect", "reason"]

    for tag in tags:
        text = re.sub(rf"<{tag}>(.|\s)*?</{tag}>", "", text, flags=re.IGNORECASE)
        text = re.sub(rf"</?{tag}>", "", text, flags=re.IGNORECASE)

    text = re.sub(r"<\|[^>]+\|>", "", text)
    return "\n".join(line for line in text.splitlines() if line.strip()).strip()

# ---------- Парсеры файлов ----------
def extract_text_from_pdf(path):
    text = ""
    reader = PdfReader(path)
    for page in reader.pages:
        if page.extract_text():
            text += page.extract_text() + "\n"
    return text.strip()

def extract_text_from_docx(path):
    texts = []
    try:
        doc = Document(path)
        for p in doc.paragraphs:
            if p.text.strip():
                texts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    texts.append(" | ".join(row_text))
        content = "\n".join(texts).strip()
        if content:
            return content
    except Exception:
        pass
    try:
        text = textract.process(path)
        return text.decode("utf-8", errors="ignore")
    except Exception:
        return ""
    
def extract_text_from_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def extract_text_from_doc_or_rtf(path):
    try:
        text = textract.process(path)
        return text.decode("utf-8", errors="ignore")
    except Exception:
        return ""

def extract_text(file_path):
    ext = file_path.lower()
    if ext.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif ext.endswith(".docx"):
        return extract_text_from_docx(file_path)
    elif ext.endswith(".txt"):
        return extract_text_from_txt(file_path)
    elif ext.endswith(".doc") or ext.endswith(".rtf"):
        return extract_text_from_doc_or_rtf(file_path)
    else:
        return ""

# ---------- Загрузка требований вакансии ----------
def load_job_requirements(path):
    if not path or not os.path.exists(path):
        return []

    text = extract_text(path)
    lines = text.splitlines()

    requirements = []
    capture = False
    for line in lines:
        if "Требования" in line:
            capture = True
            continue
        if capture:
            if not line.strip():
                break
            requirements.append(line.strip())
    return [r for r in requirements if r]

JOB_REQUIREMENTS = load_job_requirements(JOB_DESCRIPTION_PATH)

# ---------- Инструкция для модели ----------
PROMPT_INSTRUCTION = """
Ты HR-эксперт. Тебе даётся текст резюме кандидата и список требований из вакансии.
Нужно оценить соответствие кандидата каждому требованию по шкале от 1 до 10 и дать комментарий.
В конце выдай общий вывод о кандидате и интегральную оценку (среднее значение по всем требованиям, округли до десятых).

Ответ строго в формате JSON:

{
  "Оценка требований": {
    "Текст требования": {"балл": 1-10, "комментарий": "текст"},
    "Текст требования": {"балл": 1-10, "комментарий": "текст"}
  },
  "Средняя оценка": 1-10,
  "Общий вывод": "текстовое заключение"
}

⚠️ Важно: никаких комментариев вне JSON. Ключами должны быть именно тексты требований.
"""

# ---------- Основная логика ----------
def evaluate_resume(file_path):
    if not file_path:
        return "Загрузите резюме."

    resume_text = extract_text(file_path)
    if not resume_text.strip():
        return "❌ Не удалось извлечь текст из файла."

    prompt = f"""
{PROMPT_INSTRUCTION}

Требования вакансии:
{json.dumps(JOB_REQUIREMENTS, ensure_ascii=False, indent=2)}

Резюме кандидата:
{resume_text}
"""

    payload = {"model": MODEL_NAME, "messages": [{"role": "user", "content": prompt}]}

    try:
        url = f"{API_BASE_URL}/chat/completions"
        resp = requests.post(url, json=payload)
        resp.raise_for_status()
        data = resp.json()
        raw_answer = data["choices"][0]["message"]["content"]

        clean_answer = strip_thinking_tags(raw_answer)

        try:
            parsed = json.loads(clean_answer)

            result = "## 📊 Оценка по требованиям\n\n"
            scores = []
            if "Оценка требований" in parsed:
                for req, val in parsed["Оценка требований"].items():
                    score = val.get("балл", "?")
                    comment = val.get("комментарий", "")
                    result += f"**{req}:** {score} / 10 — {comment}\n\n"
                    try:
                        scores.append(int(score))
                    except Exception:
                        pass

            # средняя оценка
            if scores:
                avg_score = round(sum(scores) / len(scores), 1)
                result += f"## 📈 Средняя оценка: {avg_score} / 10\n\n"

            if "Общий вывод" in parsed:
                result += f"## 📝 Общий вывод\n{parsed['Общий вывод']}"

            return result.strip()
        except Exception:
            return f"⚠️ Ошибка парсинга JSON. Ответ модели:\n\n{clean_answer}"

    except Exception as e:
        return f"Ошибка при обращении к API: {e}"

# ---------- Интерфейс Gradio ----------
with gr.Blocks() as demo:
    gr.Markdown("## 🧑‍💼 Анализ резюме и соответствие требованиям вакансии")

    with gr.Row():
        with gr.Column():
            file_input = gr.File(
                label="Загрузите резюме кандидата",
                type="filepath",
                file_types=[".pdf", ".doc", ".docx", ".txt", ".rtf"]
            )
            submit_btn = gr.Button("Анализировать")

        with gr.Column():
            output_box = gr.Markdown(label="Результат")

    submit_btn.click(evaluate_resume, inputs=file_input, outputs=output_box)

if __name__ == "__main__":
    demo.launch()
