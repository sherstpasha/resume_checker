# utils.py
import os
import re
import json
from docx import Document
import textract
from pypdf import PdfReader
import html

try:
    from striprtf.striprtf import rtf_to_text
except Exception:
    rtf_to_text = None


SUPPORTED_REQ_EXTS = {".pdf", ".docx", ".txt", ".doc", ".rtf", ".odt"}


def strip_thinking_tags(text: str, tags=None) -> str:
    final_marker = "<|channel|>final<|message|>"
    if final_marker in text:
        text = text.split(final_marker, 1)[1]
    if tags is None:
        tags = ["think", "reflect", "reason"]
    for tag in tags:
        text = re.sub(rf"<{tag}>(.|\s)*?</{tag}>", "", text, flags=re.IGNORECASE)
        text = re.sub(rf"</?{tag}>", "", text, flags=re.IGNORECASE)
    text = re.sub(r"<\|[^>]+\|>", "", text)
    return "\n".join(line for line in text.splitlines() if line.strip()).strip()


# ---------- Парсеры файлов ----------
def extract_text_from_pdf(path):
    text = ""
    reader = PdfReader(path)
    for page in reader.pages:
        if page.extract_text():
            text += page.extract_text() + "\n"
    return text.strip()


def extract_text_from_docx(path):
    texts = []
    try:
        doc = Document(path)
        for p in doc.paragraphs:
            if p.text.strip():
                texts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_text:
                    texts.append(" | ".join(row_text))
        content = "\n".join(texts).strip()
        if content:
            return content
    except Exception:
        pass
    try:
        text = textract.process(path)
        return text.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def extract_text_from_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text_from_rtf(path):
    if rtf_to_text is not None:
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                data = f.read()
            return rtf_to_text(data) or ""
        except Exception:
            pass
    try:
        text = textract.process(path)
        return text.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def extract_text_from_doc(path):
    try:
        text = textract.process(path)
        return text.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def extract_text(file_path):
    ext = file_path.lower()
    if ext.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif ext.endswith(".docx"):
        return extract_text_from_docx(file_path)
    elif ext.endswith(".txt"):
        return extract_text_from_txt(file_path)
    elif ext.endswith(".rtf"):
        return extract_text_from_rtf(file_path)
    elif ext.endswith(".doc"):
        return extract_text_from_doc(file_path)
    else:
        return ""


# ---------- Загрузка требований ----------
def load_job_requirements(path):
    if not path or not os.path.exists(path):
        return []

    text = extract_text(path)
    lines = text.splitlines()

    requirements = []
    capture = False
    for line in lines:
        if "Требования" in line:
            capture = True
            continue
        if capture:
            if not line.strip():
                break
            requirements.append(line.strip())
    return [r for r in requirements if r]


def parse_job_paths_env(value: str | None) -> list[str]:
    if not value:
        return []
    parts = re.split(r"[,\n;]+", value)
    return [p.strip() for p in parts if p.strip()]


def load_job_requirements_many(paths: list[str]) -> list[dict]:
    """
    Возвращает список вакансий:
    [{"id": 1, "name": "<имя из файла>", "path": "<путь>", "requirements": [...]}, ...]
    Пустые/невалидные файлы игнорируются.
    """
    res = []
    for p in paths or []:
        try:
            reqs = load_job_requirements(p)
        except Exception:
            reqs = []
        if reqs:
            base = os.path.basename(p)
            name = os.path.splitext(base)[0]
            res.append(
                {"id": len(res) + 1, "name": name, "path": p, "requirements": reqs}
            )
    return res


def list_requirement_files(dir_path: str | None) -> list[str]:
    """Возвращает список файлов требований в папке (отфильтрованы по расширениям)."""
    if not dir_path:
        return []
    if not os.path.isdir(dir_path):
        return []
    files = []
    for name in os.listdir(dir_path):
        full = os.path.join(dir_path, name)
        if not os.path.isfile(full):
            continue
        _, ext = os.path.splitext(name)
        if ext.lower() in SUPPORTED_REQ_EXTS:
            files.append(full)
    return sorted(files)


def extract_json_block(text: str) -> str:
    """
    Возвращает строку с чистым JSON:
    - убирает обёртку ```json ... ``` или ```
    - извлекает первую { ... } или [ ... ]
    """
    if not text:
        return ""
    t = strip_thinking_tags(text).strip()

    t = re.sub(r"^```(?:json)?\s*|\s*```$", "", t, flags=re.IGNORECASE)

    m = re.search(r"\{.*\}", t, flags=re.DOTALL)
    if m:
        return m.group(0)
    m = re.search(r"\[.*\]", t, flags=re.DOTALL)
    if m:
        return m.group(0)
    return t.strip()


def parse_json_safely(text: str):
    """
    Пытается распарсить JSON-ответ модели, даже если он в код-блоке.
    Возвращает dict/list или бросает ValueError.
    """
    payload = extract_json_block(text)
    try:
        return json.loads(payload)
    except json.JSONDecodeError as e:
        raise ValueError(f"JSON decode failed: {e}\nExtracted: {payload[:500]}")


# ---------- Рендер и утилиты для аннотаций/резюме ----------
def parse_line_ranges_spec(spec) -> list[int]:
    """Парсер спецификаций строк: int, "1-3,5,7-8" или list[int]."""
    out: list[int] = []
    if isinstance(spec, int):
        return [spec]
    if isinstance(spec, list):
        for x in spec:
            if isinstance(x, int):
                out.append(x)
            elif isinstance(x, str) and x.isdigit():
                out.append(int(x))
        return out
    if isinstance(spec, str):
        parts = [p.strip() for p in spec.split(",") if p.strip()]
        for p in parts:
            if "-" in p:
                try:
                    s, e = p.split("-", 1)
                    s = int(s.strip())
                    e = int(e.strip())
                    if s <= e:
                        out.extend(list(range(s, e + 1)))
                except Exception:
                    continue
            elif p.isdigit():
                out.append(int(p))
    return out


def build_highlighted_resume(text: str, annotations: list[dict] | None) -> str:
    """Возвращает HTML с подсветками по строкам/фрагментам на основе аннотаций."""
    if not text:
        return '<div class="muted">Текст резюме отсутствует</div>'
    if not annotations:
        lines = text.splitlines()
        blocks = []
        for i, line_text in enumerate(lines, start=1):
            blocks.append(
                f'<div class="line" data-line="{i}">{html.escape(line_text)}</div>'
            )
        return f"<div class=\"resume-annotated\">{''.join(blocks)}</div>"

    has_line = any(
        isinstance(a, dict)
        and ("line" in a or "строка" in a or "line_from" in a or "line_to" in a or "lines" in a)
        for a in annotations
    )

    if has_line:
        lines = text.splitlines()
        per_line = {i + 1: [] for i in range(len(lines))}

        for a in annotations:
            try:
                tone = (a.get("тон") or "").lower()
                if tone not in ("good", "bad", "warn"):
                    continue

                indices: list[int] = []
                if "line" in a:
                    indices = parse_line_ranges_spec(a.get("line"))
                elif "строка" in a:
                    indices = parse_line_ranges_spec(a.get("строка"))
                elif "line_from" in a or "line_to" in a:
                    try:
                        s = int(a.get("line_from"))
                        e = int(a.get("line_to", s))
                        indices = list(range(s, e + 1)) if s <= e else []
                    except Exception:
                        indices = []
                elif "lines" in a:
                    indices = parse_line_ranges_spec(a.get("lines"))

                if not indices:
                    continue
                indices = [i for i in indices if 1 <= i <= len(lines)]
                if not indices:
                    continue

                q = a.get("цитата") or ""
                s = a.get("start")
                e = a.get("end")
                have_fragment = (s is not None and e is not None) or (q if isinstance(q, str) else "")
                comment = str(a.get("комментарий") or "").strip()
                req = str(a.get("требование") or "").strip()

                for li in indices:
                    line_text = lines[li - 1]
                    ls, le = 0, len(line_text)
                    if have_fragment:
                        ss = ee = None
                        if s is not None and e is not None:
                            try:
                                ss = int(s)
                                ee = int(e)
                            except Exception:
                                ss = ee = None
                        if not (ss is not None and ee is not None and 0 <= ss < ee <= len(line_text)):
                            if q and isinstance(q, str) and q:
                                idx = line_text.find(q)
                                if idx != -1:
                                    ss, ee = idx, idx + len(q)
                        if ss is not None and ee is not None and 0 <= ss < ee <= len(line_text):
                            ls, le = ss, ee
                    per_line[li].append(
                        {"s": ls, "e": le, "tone": tone, "comment": comment, "req": req}
                    )
            except Exception:
                continue

        blocks = []
        for i, line_text in enumerate(lines, start=1):
            spans = sorted(per_line.get(i) or [], key=lambda x: (x["s"], x["e"]))
            cleaned = []
            last_e = -1
            for sp in spans:
                if sp["s"] < last_e:
                    continue
                cleaned.append(sp)
                last_e = sp["e"]
            pos = 0
            parts = []
            for sp in cleaned:
                if pos < sp["s"]:
                    parts.append(html.escape(line_text[pos : sp["s"]]))
                tone_cls = {"good": "hl-good", "bad": "hl-bad", "warn": "hl-warn"}[sp["tone"]]
                data_comment = html.escape(sp["comment"]) if sp["comment"] else ""
                data_req = html.escape(sp["req"]) if sp["req"] else ""
                data_tone = html.escape(sp["tone"]) if sp["tone"] else ""
                content = html.escape(line_text[sp["s"] : sp["e"]])
                parts.append(
                    f'<span class="hl {tone_cls}" data-comment="{data_comment}" data-req="{data_req}" data-tone="{data_tone}">{content}</span>'
                )
                pos = sp["e"]
            if pos < len(line_text):
                parts.append(html.escape(line_text[pos:]))
            blocks.append(f"<div class=\"line\" data-line=\"{i}\">{''.join(parts)}</div>")
        return f"<div class=\"resume-annotated\">{''.join(blocks)}</div>"

    # Fallback: whole-text spans
    L = len(text)
    spans = []
    for a in annotations:
        try:
            s = int(a.get("start", -1))
            e = int(a.get("end", -1))
            tone = (a.get("тон") or "").lower()
            comment = str(a.get("комментарий") or "").strip()
            req = str(a.get("требование") or "").strip()
            quote = a.get("цитата")
            if 0 <= s < e <= L and tone in ("good", "bad", "warn"):
                spans.append({"s": s, "e": e, "tone": tone, "comment": comment, "req": req, "quote": quote or ""})
                continue
            if quote and isinstance(quote, str) and tone in ("good", "bad", "warn"):
                idx = text.find(quote)
                if idx != -1:
                    spans.append({"s": idx, "e": idx + len(quote), "tone": tone, "comment": comment, "req": req, "quote": quote})
        except Exception:
            continue
    spans.sort(key=lambda x: (x["s"], x["e"]))
    cleaned = []
    last_e = -1
    for sp in spans:
        if sp["s"] < last_e:
            continue
        cleaned.append(sp)
        last_e = sp["e"]
    out = []
    pos = 0
    for sp in cleaned:
        if pos < sp["s"]:
            out.append(html.escape(text[pos : sp["s"]]))
        tone_cls = {"good": "hl-good", "bad": "hl-bad", "warn": "hl-warn"}[sp["tone"]]
        content = html.escape(text[sp["s"] : sp["e"]])
        data_comment = html.escape(sp["comment"]) if sp["comment"] else ""
        data_req = html.escape(sp["req"]) if sp["req"] else ""
        data_tone = html.escape(sp["tone"]) if sp["tone"] else ""
        out.append(
            f'<span class="hl {tone_cls}" data-comment="{data_comment}" data-req="{data_req}" data-tone="{data_tone}">{content}</span>'
        )
        pos = sp["e"]
    if pos < L:
        out.append(html.escape(text[pos:]))
    return f"<div class=\"resume-annotated\"><pre>{''.join(out)}</pre></div>"


def city_from_address(addr: str | None) -> str:
    if not addr:
        return "—"
    s = str(addr).strip()
    for sep in [",", "—", "-", ";", "/"]:
        if sep in s:
            s = s.split(sep, 1)[0]
            break
    return s.strip() or "—"
